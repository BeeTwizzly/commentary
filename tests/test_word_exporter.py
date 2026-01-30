"""Tests for Word exporter."""

import json
import pytest
from io import BytesIO

from docx import Document

from src.export import WordExporter, ExportConfig, ExportResult, export_to_word
from src.export.formats import export_to_csv, export_to_json, export_summary_stats
from src.ui.review import ReviewSession, ReviewItem, ApprovalStatus
from src.models import HoldingData
from src.generation import GenerationResult
from src.generation.response_parser import ParsedResponse, ParsedVariation
from src.generation.llm_client import TokenUsage


@pytest.fixture
def mock_holding_contributor():
    """Create a mock contributor holding."""
    return HoldingData(
        ticker="AAPL",
        company_name="Apple Inc",
        strategy="Growth Equity",
        avg_weight=5.2,
        begin_weight=5.0,
        end_weight=5.4,
        benchmark_weight=4.0,
        benchmark_return=8.5,
        total_attribution=85.0,
        selection_effect=60.0,
        allocation_effect=25.0,
        rank=1,
        is_contributor=True,
    )


@pytest.fixture
def mock_holding_detractor():
    """Create a mock detractor holding."""
    return HoldingData(
        ticker="MSFT",
        company_name="Microsoft Corp",
        strategy="Growth Equity",
        avg_weight=3.1,
        begin_weight=3.5,
        end_weight=2.7,
        benchmark_weight=3.0,
        benchmark_return=-2.5,
        total_attribution=-42.0,
        selection_effect=-30.0,
        allocation_effect=-12.0,
        rank=1,
        is_contributor=False,
    )


@pytest.fixture
def mock_variations():
    """Create mock variations."""
    return [
        ParsedVariation(label="A", text="This is the commentary text for variation A.", word_count=9),
        ParsedVariation(label="B", text="This is variation B text.", word_count=5),
        ParsedVariation(label="C", text="Variation C.", word_count=2),
    ]


@pytest.fixture
def mock_result(mock_variations):
    """Create a mock generation result."""
    parsed = ParsedResponse(
        variations=mock_variations,
        raw_response="raw content",
        parse_warnings=[],
    )
    return GenerationResult(
        parsed=parsed,
        usage=TokenUsage(prompt_tokens=100, completion_tokens=50, total_tokens=150),
        cost_usd=0.01,
        latency_seconds=1.5,
        model="gpt-4o",
        success=True,
        error_message="",
    )


@pytest.fixture
def review_session(mock_holding_contributor, mock_holding_detractor, mock_result):
    """Create a review session with approved items."""
    item1 = ReviewItem(
        holding=mock_holding_contributor,
        strategy="Growth Equity",
        result=mock_result,
    )
    item1.approve()

    item2 = ReviewItem(
        holding=mock_holding_detractor,
        strategy="Growth Equity",
        result=mock_result,
    )
    item2.approve()

    return ReviewSession(
        strategy="Growth Equity",
        quarter="Q4 2025",
        items=[item1, item2],
    )


@pytest.fixture
def empty_session():
    """Create session with no approved items."""
    return ReviewSession(
        strategy="Growth",
        quarter="Q4 2025",
        items=[],
    )


class TestExportConfig:
    """Tests for ExportConfig."""

    def test_default_config(self):
        """Test default configuration values."""
        config = ExportConfig()

        assert config.include_metadata is True
        assert config.include_statistics is True
        assert config.group_by_type is True
        assert config.include_effects is True
        assert config.font_name == "Calibri"
        assert config.font_size == 11

    def test_custom_config(self):
        """Test custom configuration."""
        config = ExportConfig(
            include_metadata=False,
            font_name="Arial",
            font_size=10,
        )

        assert config.include_metadata is False
        assert config.font_name == "Arial"


class TestWordExporter:
    """Tests for WordExporter class."""

    def test_export_creates_document(self, review_session):
        """Test that export creates a valid Word document."""
        exporter = WordExporter()
        result = exporter.export(review_session)

        assert result.success is True
        assert result.document is not None
        assert result.buffer is not None
        assert result.item_count == 2
        assert result.filename.endswith(".docx")

    def test_export_buffer_is_valid_docx(self, review_session):
        """Test that the buffer contains a valid docx."""
        exporter = WordExporter()
        result = exporter.export(review_session)

        # Try to read the buffer as a Document
        doc = Document(result.buffer)

        # Check it has content
        assert len(doc.paragraphs) > 0

    def test_export_includes_strategy_and_quarter(self, review_session):
        """Test that document includes strategy and quarter."""
        exporter = WordExporter()
        result = exporter.export(review_session)

        doc = Document(result.buffer)
        text = "\n".join(p.text for p in doc.paragraphs)

        assert "Growth Equity" in text
        assert "Q4 2025" in text

    def test_export_includes_holdings(self, review_session):
        """Test that document includes holding names."""
        exporter = WordExporter()
        result = exporter.export(review_session)

        doc = Document(result.buffer)
        text = "\n".join(p.text for p in doc.paragraphs)

        assert "Apple" in text
        assert "AAPL" in text
        assert "Microsoft" in text
        assert "MSFT" in text

    def test_export_groups_by_type(self, review_session):
        """Test that items are grouped by contributor/detractor."""
        config = ExportConfig(group_by_type=True)
        exporter = WordExporter(config)
        result = exporter.export(review_session)

        doc = Document(result.buffer)
        text = "\n".join(p.text for p in doc.paragraphs)

        assert "Contributors" in text
        assert "Detractors" in text

    def test_export_includes_effects(self, review_session):
        """Test that attribution effects are included."""
        config = ExportConfig(include_effects=True)
        exporter = WordExporter(config)
        result = exporter.export(review_session)

        doc = Document(result.buffer)
        text = "\n".join(p.text for p in doc.paragraphs)

        assert "+85.0 bps" in text
        assert "-42.0 bps" in text

    def test_export_excludes_effects_when_disabled(self, review_session):
        """Test effects are excluded when disabled."""
        config = ExportConfig(include_effects=False)
        exporter = WordExporter(config)
        result = exporter.export(review_session)

        doc = Document(result.buffer)
        text = "\n".join(p.text for p in doc.paragraphs)

        assert "+85.0 bps" not in text
        assert "-42.0 bps" not in text

    def test_export_no_approved_items(self, mock_holding_contributor, mock_result):
        """Test export with no approved items fails gracefully."""
        item = ReviewItem(
            holding=mock_holding_contributor,
            strategy="Growth",
            result=mock_result,
        )
        # Don't approve

        session = ReviewSession(
            strategy="Growth",
            quarter="Q4 2025",
            items=[item],
        )

        exporter = WordExporter()
        result = exporter.export(session)

        assert result.success is False
        assert "No approved items" in result.error_message

    def test_export_to_text(self, review_session):
        """Test plain text export."""
        exporter = WordExporter()
        text = exporter.export_to_text(review_session)

        assert "Growth Equity" in text
        assert "Q4 2025" in text
        assert "AAPL" in text
        assert "MSFT" in text
        assert "CONTRIBUTORS" in text
        assert "DETRACTORS" in text

    def test_filename_format(self, review_session):
        """Test that filename is properly formatted."""
        exporter = WordExporter()
        result = exporter.export(review_session)

        assert "Growth" in result.filename
        assert "Q4_2025" in result.filename
        assert ".docx" in result.filename


class TestExportFormats:
    """Tests for additional export formats."""

    def test_export_to_csv(self, review_session):
        """Test CSV export."""
        csv_content = export_to_csv(review_session)

        # Check header
        assert "Ticker" in csv_content
        assert "Company" in csv_content
        assert "Commentary" in csv_content

        # Check data
        assert "AAPL" in csv_content
        assert "Apple Inc" in csv_content
        assert "Contributor" in csv_content
        assert "Detractor" in csv_content

    def test_export_to_json(self, review_session):
        """Test JSON export."""
        json_content = export_to_json(review_session)
        data = json.loads(json_content)

        assert "metadata" in data
        assert "commentary" in data
        assert data["metadata"]["strategy"] == "Growth Equity"
        assert data["metadata"]["quarter"] == "Q4 2025"
        assert data["metadata"]["total_items"] == 2
        assert len(data["commentary"]) == 2

    def test_export_summary_stats(self, review_session):
        """Test summary statistics generation."""
        stats = export_summary_stats(review_session)

        assert stats["strategy"] == "Growth Equity"
        assert stats["quarter"] == "Q4 2025"
        assert stats["total_approved"] == 2
        assert stats["contributors_approved"] == 1
        assert stats["detractors_approved"] == 1
        assert "avg_word_count" in stats
        assert "total_cost_usd" in stats

    def test_export_summary_stats_empty(self, empty_session):
        """Test summary stats with no items."""
        stats = export_summary_stats(empty_session)

        assert "error" in stats


class TestExportToWordConvenience:
    """Tests for convenience function."""

    def test_export_to_word_function(self, review_session):
        """Test the convenience function."""
        result = export_to_word(review_session)

        assert result.success is True
        assert result.item_count == 2

    def test_export_to_word_with_config(self, review_session):
        """Test convenience function with custom config."""
        config = ExportConfig(include_statistics=False)
        result = export_to_word(review_session, config)

        assert result.success is True
