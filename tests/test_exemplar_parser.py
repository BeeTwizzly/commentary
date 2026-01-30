"""Tests for the exemplar parser and selector modules."""

import pytest
import json
from pathlib import Path

from docx import Document

from src.parsers.exemplar_parser import (
    parse_word_document,
    parse_directory,
    save_exemplars_json,
    load_exemplars_json,
    _parse_filename_metadata,
    _detect_section_header,
)
from src.parsers.exemplar_selector import ExemplarSelector
from src.models import ExemplarBlurb, ExemplarSelection


@pytest.fixture
def sample_docx(tmp_path) -> Path:
    """Create a sample Word document with commentary blurbs."""
    doc = Document()

    # Add some intro text
    doc.add_paragraph("PORTFOLIO PERFORMANCE AND ATTRIBUTION")
    doc.add_paragraph("")

    # Contributors section
    doc.add_paragraph("Contributors")
    doc.add_paragraph("")

    doc.add_paragraph(
        "Apple Inc. (AAPL): Apple continued to deliver exceptional results as Services "
        "revenue grew 14% year-over-year, demonstrating the durability of its ecosystem "
        "monetization strategy. The installed base expansion in emerging markets provides "
        "runway for continued growth. Management's capital allocation remains shareholder-friendly. "
        "We maintain our position given attractive risk-adjusted return potential."
    )

    doc.add_paragraph(
        "NVIDIA Corporation (NVDA): Our largest active overweight contributed meaningfully "
        "as datacenter revenue exceeded expectations by 15%. Management's commentary on "
        "enterprise AI adoption reinforced our thesis on sustainable demand. The different "
        "risk profile at current scale warrants monitoring but fundamentals remain intact."
    )

    doc.add_paragraph("")

    # Detractors section
    doc.add_paragraph("Detractors")
    doc.add_paragraph("")

    doc.add_paragraph(
        "Intel Corporation (INTC): Intel detracted from performance as foundry delays "
        "continued to push out the roadmap for process leadership. While valuation appears "
        "undemanding, execution risk remains elevated. We have trimmed the position pending "
        "evidence of manufacturing progress. The turnaround thesis requires patience."
    )

    doc.add_paragraph(
        "PayPal Holdings Inc. (PYPL): PayPal weighed on returns as competitive pressures "
        "in the digital payments space intensified. Management's strategic pivot toward "
        "branded checkout shows early promise but will take time to materialize. We are "
        "monitoring closely for signs of stabilization in take rates."
    )

    file_path = tmp_path / "Q3_2025_Commentary.docx"
    doc.save(file_path)
    return file_path


@pytest.fixture
def sample_docx_multiline(tmp_path) -> Path:
    """Create a document with header and text on separate lines."""
    doc = Document()

    doc.add_paragraph("Contributors")
    doc.add_paragraph("")

    # Header on its own line
    doc.add_paragraph("Microsoft Corporation (MSFT):")
    doc.add_paragraph(
        "Microsoft delivered strong results driven by Azure growth exceeding 30% year-over-year. "
        "The AI integration across the product suite positions the company well for enterprise "
        "adoption. Cloud margins continue to expand as scale benefits accrue. We remain "
        "constructive on the long-term opportunity."
    )

    file_path = tmp_path / "Q2_2025_Report.docx"
    doc.save(file_path)
    return file_path


@pytest.fixture
def sample_exemplars_json(tmp_path) -> Path:
    """Create a sample exemplars JSON file."""
    data = {
        "metadata": {
            "total_blurbs": 4,
            "unique_tickers": 4,
            "contributors": 2,
            "detractors": 2,
        },
        "blurbs_by_ticker": {
            "AAPL": [{
                "ticker": "AAPL",
                "company_name": "Apple Inc.",
                "blurb_text": "Apple delivered strong results with Services growth of 14%. "
                              "The ecosystem monetization continues to demonstrate durability. "
                              "Capital allocation remains shareholder friendly. We maintain position.",
                "quarter": "Q3",
                "year": 2025,
                "blurb_type": "contributor",
                "word_count": 45,
                "source_file": "Q3_2025.docx",
            }],
            "NVDA": [{
                "ticker": "NVDA",
                "company_name": "NVIDIA Corporation",
                "blurb_text": "NVIDIA contributed meaningfully as datacenter revenue beat expectations. "
                              "Management commentary reinforced sustainable AI demand thesis. "
                              "Fundamentals remain intact despite different risk profile.",
                "quarter": "Q3",
                "year": 2025,
                "blurb_type": "contributor",
                "word_count": 42,
                "source_file": "Q3_2025.docx",
            }],
            "INTC": [{
                "ticker": "INTC",
                "company_name": "Intel Corporation",
                "blurb_text": "Intel detracted as foundry delays pushed out the manufacturing roadmap. "
                              "Valuation undemanding but execution risk elevated. "
                              "Position trimmed pending progress evidence.",
                "quarter": "Q3",
                "year": 2025,
                "blurb_type": "detractor",
                "word_count": 38,
                "source_file": "Q3_2025.docx",
            }],
            "PYPL": [{
                "ticker": "PYPL",
                "company_name": "PayPal Holdings Inc.",
                "blurb_text": "PayPal weighed on returns amid competitive pressures in digital payments. "
                              "Strategic pivot shows early promise but requires time. "
                              "Monitoring for take rate stabilization.",
                "quarter": "Q3",
                "year": 2025,
                "blurb_type": "detractor",
                "word_count": 40,
                "source_file": "Q3_2025.docx",
            }],
        },
    }

    json_path = tmp_path / "exemplars.json"
    with json_path.open("w") as f:
        json.dump(data, f)

    return json_path


class TestParseWordDocument:
    """Tests for Word document parsing."""

    def test_extracts_blurbs(self, sample_docx):
        """Should extract all blurbs from document."""
        blurbs = parse_word_document(sample_docx)

        assert len(blurbs) >= 4
        tickers = {b.ticker for b in blurbs}
        assert tickers == {"AAPL", "NVDA", "INTC", "PYPL"}

    def test_identifies_contributor_type(self, sample_docx):
        """Should correctly identify contributor blurbs."""
        blurbs = parse_word_document(sample_docx)

        aapl = next(b for b in blurbs if b.ticker == "AAPL")
        assert aapl.blurb_type == "contributor"

        nvda = next(b for b in blurbs if b.ticker == "NVDA")
        assert nvda.blurb_type == "contributor"

    def test_identifies_detractor_type(self, sample_docx):
        """Should correctly identify detractor blurbs."""
        blurbs = parse_word_document(sample_docx)

        intc = next(b for b in blurbs if b.ticker == "INTC")
        assert intc.blurb_type == "detractor"

        pypl = next(b for b in blurbs if b.ticker == "PYPL")
        assert pypl.blurb_type == "detractor"

    def test_extracts_company_name(self, sample_docx):
        """Should extract company name correctly."""
        blurbs = parse_word_document(sample_docx)

        aapl = next(b for b in blurbs if b.ticker == "AAPL")
        assert aapl.company_name == "Apple Inc."

    def test_calculates_word_count(self, sample_docx):
        """Should calculate word count for each blurb."""
        blurbs = parse_word_document(sample_docx)

        for blurb in blurbs:
            assert blurb.word_count > 0
            assert blurb.word_count == len(blurb.blurb_text.split())

    def test_extracts_quarter_from_filename(self, sample_docx):
        """Should extract quarter info from filename."""
        blurbs = parse_word_document(sample_docx)

        assert all(b.quarter == "Q3" for b in blurbs)
        assert all(b.year == 2025 for b in blurbs)

    def test_handles_multiline_format(self, sample_docx_multiline):
        """Should handle blurbs split across paragraphs."""
        blurbs = parse_word_document(sample_docx_multiline)

        assert len(blurbs) >= 1
        msft = next((b for b in blurbs if b.ticker == "MSFT"), None)
        assert msft is not None
        assert "Azure" in msft.blurb_text

    def test_raises_on_missing_file(self, tmp_path):
        """Should raise FileNotFoundError for missing file."""
        with pytest.raises(FileNotFoundError):
            parse_word_document(tmp_path / "nonexistent.docx")


class TestParseFilenameMetadata:
    """Tests for filename metadata extraction."""

    def test_q3_2025_format(self):
        assert _parse_filename_metadata("Q3_2025_Commentary.docx") == ("Q3", 2025)

    def test_2025_q3_format(self):
        assert _parse_filename_metadata("2025_Q3_Report.docx") == ("Q3", 2025)

    def test_3q25_format(self):
        assert _parse_filename_metadata("3Q25_Report.docx") == ("Q3", 2025)

    def test_returns_default_for_unknown(self):
        quarter, year = _parse_filename_metadata("random_file.docx")
        assert quarter == "Q4"
        assert year == 2025


class TestDetectSectionHeader:
    """Tests for section header detection."""

    def test_detects_contributors(self):
        assert _detect_section_header("Contributors") == "contributor"
        assert _detect_section_header("Top Contributors") == "contributor"
        assert _detect_section_header("CONTRIBUTORS") == "contributor"

    def test_detects_detractors(self):
        assert _detect_section_header("Detractors") == "detractor"
        assert _detect_section_header("What Didn't Work") == "detractor"

    def test_returns_none_for_regular_text(self):
        assert _detect_section_header("Apple delivered strong results...") is None
        assert _detect_section_header("The portfolio outperformed the benchmark.") is None


class TestSaveAndLoadJson:
    """Tests for JSON serialization."""

    def test_roundtrip(self, sample_docx, tmp_path):
        """Should save and reload blurbs correctly."""
        blurbs = parse_word_document(sample_docx)
        json_path = tmp_path / "output.json"

        save_exemplars_json(blurbs, json_path)
        reloaded = load_exemplars_json(json_path)

        assert len(reloaded) == len(blurbs)

        original_tickers = {b.ticker for b in blurbs}
        reloaded_tickers = {b.ticker for b in reloaded}
        assert original_tickers == reloaded_tickers

    def test_raises_on_missing_json(self, tmp_path):
        """Should raise FileNotFoundError for missing JSON."""
        with pytest.raises(FileNotFoundError):
            load_exemplars_json(tmp_path / "nonexistent.json")


class TestExemplarSelector:
    """Tests for exemplar selection logic."""

    def test_loads_from_json(self, sample_exemplars_json):
        """Should load selector from JSON file."""
        selector = ExemplarSelector.load(sample_exemplars_json)

        assert selector.total_blurbs == 4
        assert len(selector.available_tickers) == 4

    def test_selects_same_ticker_exemplar(self, sample_exemplars_json):
        """Should include prior blurb for same ticker."""
        selector = ExemplarSelector.load(sample_exemplars_json)

        selection = selector.select("AAPL", is_contributor=True)

        assert selection.same_ticker_exemplar is not None
        assert selection.same_ticker_exemplar.ticker == "AAPL"

    def test_selects_matching_type(self, sample_exemplars_json):
        """Should select exemplars matching contributor/detractor type."""
        selector = ExemplarSelector.load(sample_exemplars_json)

        contrib_selection = selector.select("NEW", is_contributor=True, count=3)
        detract_selection = selector.select("NEW", is_contributor=False, count=3)

        for ex in contrib_selection.get_all_exemplars():
            assert ex.blurb_type == "contributor"

        for ex in detract_selection.get_all_exemplars():
            assert ex.blurb_type == "detractor"

    def test_excludes_specified_tickers(self, sample_exemplars_json):
        """Should exclude specified tickers from selection."""
        selector = ExemplarSelector.load(sample_exemplars_json)

        selection = selector.select(
            "NEW",
            is_contributor=True,
            exclude_tickers=["AAPL"]
        )

        for ex in selection.similar_exemplars:
            assert ex.ticker != "AAPL"

    def test_format_for_prompt(self, sample_exemplars_json):
        """Should format selection as prompt text."""
        selector = ExemplarSelector.load(sample_exemplars_json)

        selection = selector.select("NEW", is_contributor=True, count=2)
        prompt_text = selection.format_for_prompt()

        assert "Example 1" in prompt_text
        assert "contributor" in prompt_text.lower()

    def test_handles_unknown_ticker(self, sample_exemplars_json):
        """Should handle ticker with no prior blurbs."""
        selector = ExemplarSelector.load(sample_exemplars_json)

        selection = selector.select("UNKNOWN", is_contributor=True)

        assert selection.same_ticker_exemplar is None
        assert len(selection.similar_exemplars) > 0


class TestExemplarBlurb:
    """Tests for ExemplarBlurb model."""

    def test_matches_type_contributor(self):
        """Should correctly match contributor type."""
        blurb = ExemplarBlurb(
            ticker="TEST",
            company_name="Test Co",
            blurb_text="Test blurb",
            quarter="Q1",
            year=2025,
            blurb_type="contributor",
            word_count=10,
            source_file="test.docx",
        )

        assert blurb.matches_type(is_contributor=True) is True
        assert blurb.matches_type(is_contributor=False) is False

    def test_matches_type_detractor(self):
        """Should correctly match detractor type."""
        blurb = ExemplarBlurb(
            ticker="TEST",
            company_name="Test Co",
            blurb_text="Test blurb",
            quarter="Q1",
            year=2025,
            blurb_type="detractor",
            word_count=10,
            source_file="test.docx",
        )

        assert blurb.matches_type(is_contributor=False) is True
        assert blurb.matches_type(is_contributor=True) is False
