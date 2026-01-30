"""Word document exporter for portfolio commentary."""

from __future__ import annotations

import logging
from dataclasses import dataclass
from datetime import datetime
from io import BytesIO

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

from src.ui.review import ReviewSession, ReviewItem

logger = logging.getLogger(__name__)


@dataclass
class ExportConfig:
    """
    Configuration for Word export.

    Attributes:
        include_metadata: Include header with quarter, date, stats
        include_statistics: Include generation statistics section
        group_by_type: Group by contributors/detractors within strategy
        include_effects: Include attribution effect values
        font_name: Font family for body text
        font_size: Font size in points
        heading_font_size: Font size for headings
        company_format: Format string for company name line
    """
    include_metadata: bool = True
    include_statistics: bool = True
    group_by_type: bool = True
    include_effects: bool = True
    font_name: str = "Calibri"
    font_size: int = 11
    heading_font_size: int = 14
    company_format: str = "{company} ({ticker})"


@dataclass
class ExportResult:
    """
    Result of an export operation.

    Attributes:
        success: Whether export succeeded
        document: The generated Document object (if success)
        buffer: BytesIO buffer with document bytes (if success)
        filename: Suggested filename
        item_count: Number of items exported
        error_message: Error message (if failed)
    """
    success: bool
    document: Document | None = None
    buffer: BytesIO | None = None
    filename: str = ""
    item_count: int = 0
    error_message: str = ""


class WordExporter:
    """
    Exports approved commentary to Word documents.

    Usage:
        exporter = WordExporter(config)
        result = exporter.export(session)

        if result.success:
            # result.buffer contains the .docx bytes
            # result.filename has suggested name
    """

    def __init__(self, config: ExportConfig | None = None):
        """
        Initialize exporter with configuration.

        Args:
            config: Export configuration (uses defaults if None)
        """
        self.config = config or ExportConfig()

    def export(self, session: ReviewSession) -> ExportResult:
        """
        Export approved items from a review session to Word.

        Args:
            session: Review session with approved items

        Returns:
            ExportResult with document buffer or error
        """
        try:
            # Get approved items
            items = session.exportable_items

            if not items:
                return ExportResult(
                    success=False,
                    error_message="No approved items to export",
                )

            # Create document
            doc = Document()
            self._setup_styles(doc)

            # Add content
            if self.config.include_metadata:
                self._add_metadata_header(doc, session)

            self._add_commentary_sections(doc, session, items)

            if self.config.include_statistics:
                self._add_statistics(doc, session, items)

            # Generate buffer
            buffer = BytesIO()
            doc.save(buffer)
            buffer.seek(0)

            # Generate filename
            date_str = datetime.now().strftime("%Y%m%d")
            quarter_clean = session.quarter.replace(" ", "_")
            strategy_clean = "".join(c if c.isalnum() else "_" for c in session.strategy)
            filename = f"{strategy_clean}_{quarter_clean}_Commentary_{date_str}.docx"

            logger.info(
                "Exported %d items to Word document: %s",
                len(items),
                filename,
            )

            return ExportResult(
                success=True,
                document=doc,
                buffer=buffer,
                filename=filename,
                item_count=len(items),
            )

        except Exception as e:
            logger.exception("Export failed")
            return ExportResult(
                success=False,
                error_message=str(e),
            )

    def _setup_styles(self, doc: Document) -> None:
        """Configure document styles."""
        # Set default font
        style = doc.styles["Normal"]
        font = style.font
        font.name = self.config.font_name
        font.size = Pt(self.config.font_size)

        # Heading 1 style
        if "Heading 1" in doc.styles:
            h1 = doc.styles["Heading 1"]
            h1.font.name = self.config.font_name
            h1.font.size = Pt(self.config.heading_font_size)
            h1.font.bold = True
            h1.font.color.rgb = RGBColor(0x1E, 0x3A, 0x5F)  # Dark blue

        # Heading 2 style
        if "Heading 2" in doc.styles:
            h2 = doc.styles["Heading 2"]
            h2.font.name = self.config.font_name
            h2.font.size = Pt(self.config.heading_font_size - 2)
            h2.font.bold = True
            h2.font.color.rgb = RGBColor(0x2E, 0x5A, 0x7F)

    def _add_metadata_header(self, doc: Document, session: ReviewSession) -> None:
        """Add metadata header to document."""
        # Title
        title = doc.add_heading(f"{session.strategy} Portfolio Commentary", level=0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Subtitle with quarter
        subtitle = doc.add_paragraph()
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = subtitle.add_run(session.quarter)
        run.font.size = Pt(14)
        run.font.italic = True

        # Generation info
        info = doc.add_paragraph()
        info.alignment = WD_ALIGN_PARAGRAPH.CENTER
        info_run = info.add_run(
            f"Generated: {datetime.now().strftime('%B %d, %Y at %I:%M %p')}"
        )
        info_run.font.size = Pt(10)
        info_run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

        doc.add_paragraph()  # Spacer

    def _add_commentary_sections(
        self,
        doc: Document,
        session: ReviewSession,
        items: list[ReviewItem],
    ) -> None:
        """Add commentary content organized by type."""
        if self.config.group_by_type:
            # Separate contributors and detractors
            contributors = [i for i in items if i.is_contributor]
            detractors = [i for i in items if not i.is_contributor]

            # Sort by absolute effect (highest first)
            contributors.sort(key=lambda x: abs(x.holding.total_attribution or 0), reverse=True)
            detractors.sort(key=lambda x: abs(x.holding.total_attribution or 0), reverse=True)

            if contributors:
                doc.add_heading("Top Contributors", level=1)
                self._add_items_section(doc, contributors)

            if detractors:
                doc.add_heading("Top Detractors", level=1)
                self._add_items_section(doc, detractors)
        else:
            # All items together
            doc.add_heading("Commentary", level=1)
            self._add_items_section(doc, items)

    def _add_items_section(self, doc: Document, items: list[ReviewItem]) -> None:
        """Add a section of commentary items."""
        for item in items:
            # Company name heading
            company_line = self.config.company_format.format(
                company=item.company_name,
                ticker=item.ticker,
            )

            if self.config.include_effects and item.holding.total_attribution is not None:
                effect = item.holding.total_attribution
                effect_str = f" [{effect:+.1f} bps]"
                company_line += effect_str

            # Add as bold paragraph (not heading to keep it compact)
            p = doc.add_paragraph()
            run = p.add_run(company_line)
            run.bold = True
            run.font.size = Pt(self.config.font_size + 1)

            # Commentary text
            text = item.current_text
            if text:
                commentary_p = doc.add_paragraph(text)
                commentary_p.paragraph_format.space_after = Pt(12)
            else:
                doc.add_paragraph("[No commentary text]")

    def _add_statistics(
        self,
        doc: Document,
        session: ReviewSession,
        items: list[ReviewItem],
    ) -> None:
        """Add statistics section at end of document."""
        doc.add_page_break()
        doc.add_heading("Export Statistics", level=1)

        # Calculate stats
        word_counts = [i.current_word_count for i in items]
        avg_words = sum(word_counts) / len(word_counts) if word_counts else 0
        total_cost = sum(i.result.cost_usd for i in items if i.result)

        # Create stats table
        table = doc.add_table(rows=7, cols=2)
        table.style = "Table Grid"
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        stats = [
            ("Strategy", session.strategy),
            ("Quarter", session.quarter),
            ("Holdings Exported", str(len(items))),
            ("Contributors", str(sum(1 for i in items if i.is_contributor))),
            ("Detractors", str(sum(1 for i in items if not i.is_contributor))),
            ("Average Word Count", f"{avg_words:.0f}"),
            ("Total Generation Cost", f"${total_cost:.4f}"),
        ]

        for i, (label, value) in enumerate(stats):
            row = table.rows[i]
            row.cells[0].text = label
            row.cells[1].text = value

            # Bold the label
            for paragraph in row.cells[0].paragraphs:
                for run in paragraph.runs:
                    run.bold = True

    def export_to_text(self, session: ReviewSession) -> str:
        """
        Export approved items as plain text (clipboard-friendly).

        Args:
            session: Review session with approved items

        Returns:
            Plain text string with all commentary
        """
        items = session.exportable_items

        if not items:
            return ""

        lines = []

        # Header
        lines.append(f"{session.strategy} Portfolio Commentary")
        lines.append(f"{session.quarter}")
        lines.append("=" * 50)
        lines.append("")

        if self.config.group_by_type:
            contributors = [i for i in items if i.is_contributor]
            detractors = [i for i in items if not i.is_contributor]

            contributors.sort(key=lambda x: abs(x.holding.total_attribution or 0), reverse=True)
            detractors.sort(key=lambda x: abs(x.holding.total_attribution or 0), reverse=True)

            if contributors:
                lines.append("TOP CONTRIBUTORS")
                lines.append("-" * 30)
                for item in contributors:
                    lines.extend(self._format_item_text(item))
                lines.append("")

            if detractors:
                lines.append("TOP DETRACTORS")
                lines.append("-" * 30)
                for item in detractors:
                    lines.extend(self._format_item_text(item))
        else:
            for item in items:
                lines.extend(self._format_item_text(item))

        return "\n".join(lines)

    def _format_item_text(self, item: ReviewItem) -> list[str]:
        """Format a single item as text lines."""
        lines = []

        header = f"{item.company_name} ({item.ticker})"
        if self.config.include_effects and item.holding.total_attribution:
            header += f" [{item.holding.total_attribution:+.1f} bps]"

        lines.append(header)
        lines.append(item.current_text)
        lines.append("")

        return lines


def export_to_word(session: ReviewSession, config: ExportConfig | None = None) -> ExportResult:
    """
    Convenience function to export a session to Word.

    Args:
        session: Review session with approved items
        config: Export configuration (optional)

    Returns:
        ExportResult with document buffer
    """
    exporter = WordExporter(config)
    return exporter.export(session)
